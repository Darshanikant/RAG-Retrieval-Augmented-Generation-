# loader.py
import os
from typing import List
from PyPDF2 import PdfReader
import docx
from langchain.schema import Document


def load_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def load_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def load_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_documents(folder_path: str) -> List[Document]:
    documents = []
    for filename in os.listdir(folder_path):
        full_path = os.path.join(folder_path, filename)
        if filename.endswith(".txt"):
            content = load_txt(full_path)
        elif filename.endswith(".pdf"):
            content = load_pdf(full_path)
        elif filename.endswith(".docx"):
            content = load_docx(full_path)
        else:
            print(f"Unsupported file format: {filename}")
            continue

        documents.append(Document(page_content=content, metadata={"source": filename}))
    return documents
